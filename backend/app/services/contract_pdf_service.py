import os
import math
import re
from datetime import date
from io import BytesIO
from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML
from docx import Document as DocxDocument
from app.services.supabase_service import supabase_service

TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "templates")

env = Environment(loader=FileSystemLoader(TEMPLATE_DIR))

STATUS_LABELS = {
    "pending": "待生效",
    "active": "生效中",
    "expired": "已過期",
    "terminated": "已終止",
}

STUDENT_DETAIL_TYPE_LABELS = {
    "lesson_price": "課程單價",
    "discount": "優惠折扣",
    "compensation": "補償堂數",
}

TEACHER_DETAIL_TYPE_LABELS = {
    "course_rate": "課程費率",
    "base_salary": "底薪",
    "allowance": "津貼",
}

EMPLOYMENT_TYPE_LABELS = {
    "hourly": "時薪",
    "full_time": "正職",
}

WEEKDAY_LABELS = {
    0: "週一",
    1: "週二",
    2: "週三",
    3: "週四",
    4: "週五",
    5: "週六",
    6: "週日",
}


async def _fetch_student_contract_data(contract_id: str) -> dict:
    """查詢學生合約 + 明細 + 關聯名稱"""
    result = await supabase_service.table_select(
        table="student_contracts",
        select="id,contract_no,student_id,contract_status,start_date,end_date,total_lessons,remaining_lessons,total_amount,total_leave_allowed,is_recurring,notes",
        filters={"id": contract_id, "is_deleted": "eq.false"},
    )
    if not result:
        return None

    contract = result[0]

    # 學生名稱
    if contract.get("student_id"):
        student = await supabase_service.table_select(
            table="students", select="name",
            filters={"id": contract["student_id"]},
        )
        contract["student_name"] = student[0]["name"] if student else None

    # 明細
    details = await supabase_service.table_select(
        table="student_contract_details",
        select="id,detail_type,course_id,description,amount",
        filters={
            "student_contract_id": f"eq.{contract_id}",
            "is_deleted": "eq.false",
        },
    )
    for d in details:
        if d.get("course_id"):
            course = await supabase_service.table_select(
                table="courses", select="course_name",
                filters={"id": d["course_id"]},
            )
            d["course_name"] = course[0]["course_name"] if course else None

    return {"contract": contract, "details": details}


async def _fetch_teacher_contract_data(contract_id: str) -> dict:
    """查詢教師合約 + 明細 + 工作時段 + 關聯名稱"""
    result = await supabase_service.table_select(
        table="teacher_contracts",
        select="id,contract_no,teacher_id,contract_status,start_date,end_date,employment_type,trial_completed_bonus,trial_to_formal_bonus,notes",
        filters={"id": contract_id, "is_deleted": "eq.false"},
    )
    if not result:
        return None

    contract = result[0]

    # 教師名稱
    if contract.get("teacher_id"):
        teacher = await supabase_service.table_select(
            table="teachers", select="name",
            filters={"id": contract["teacher_id"]},
        )
        contract["teacher_name"] = teacher[0]["name"] if teacher else None

    # 明細
    details = await supabase_service.table_select(
        table="teacher_contract_details",
        select="id,detail_type,course_id,description,amount",
        filters={
            "teacher_contract_id": f"eq.{contract_id}",
            "is_deleted": "eq.false",
        },
    )
    for d in details:
        if d.get("course_id"):
            course = await supabase_service.table_select(
                table="courses", select="course_name",
                filters={"id": d["course_id"]},
            )
            d["course_name"] = course[0]["course_name"] if course else None

    # 工作時段
    work_schedules = await supabase_service.table_select(
        table="teacher_work_schedules",
        select="id,weekday,start_time,end_time,notes",
        filters={
            "teacher_contract_id": f"eq.{contract_id}",
            "is_deleted": "eq.false",
        },
    )

    return {"contract": contract, "details": details, "work_schedules": work_schedules}


async def generate_student_contract_pdf(contract_id: str) -> tuple[bytes, str] | None:
    """產生學生合約 PDF，回傳 (pdf_bytes, contract_no) 或 None"""
    data = await _fetch_student_contract_data(contract_id)
    if not data:
        return None

    template = env.get_template("student_contract.html")
    html_str = template.render(
        contract=data["contract"],
        details=data["details"],
        status_label=STATUS_LABELS.get(data["contract"]["contract_status"], data["contract"]["contract_status"]),
        detail_type_labels=STUDENT_DETAIL_TYPE_LABELS,
    )

    pdf_bytes = HTML(string=html_str).write_pdf()
    return pdf_bytes, data["contract"]["contract_no"]


async def generate_addendum_pdf(addendum_id: str) -> tuple[bytes, str] | None:
    """產生附約 PDF，回傳 (pdf_bytes, addendum_no) 或 None"""
    result = await supabase_service.table_select(
        table="contract_addendums",
        select="id,addendum_no,contract_type,parent_contract_id,original_end_date,new_end_date,addendum_status,notes",
        filters={"id": addendum_id, "is_deleted": "eq.false"},
    )
    if not result:
        return None

    addendum = result[0]
    contract_type = addendum["contract_type"]

    # 取得母約編號和人名
    if contract_type == "student":
        parent = await supabase_service.table_select(
            table="student_contracts",
            select="contract_no,student_id",
            filters={"id": addendum["parent_contract_id"], "is_deleted": "eq.false"},
        )
        if not parent:
            return None
        parent_contract_no = parent[0]["contract_no"]
        person_name = None
        if parent[0].get("student_id"):
            student = await supabase_service.table_select(
                table="students", select="name",
                filters={"id": parent[0]["student_id"]},
            )
            person_name = student[0]["name"] if student else None
        title = "學生"
        person_label = "學生"
    else:
        parent = await supabase_service.table_select(
            table="teacher_contracts",
            select="contract_no,teacher_id",
            filters={"id": addendum["parent_contract_id"], "is_deleted": "eq.false"},
        )
        if not parent:
            return None
        parent_contract_no = parent[0]["contract_no"]
        person_name = None
        if parent[0].get("teacher_id"):
            teacher = await supabase_service.table_select(
                table="teachers", select="name",
                filters={"id": parent[0]["teacher_id"]},
            )
            person_name = teacher[0]["name"] if teacher else None
        title = "教師"
        person_label = "教師"

    template = env.get_template("contract_addendum.html")
    html_str = template.render(
        addendum=addendum,
        title=title,
        person_label=person_label,
        parent_contract_no=parent_contract_no,
        person_name=person_name,
        status_label=STATUS_LABELS.get(addendum["addendum_status"], addendum["addendum_status"]),
    )

    pdf_bytes = HTML(string=html_str).write_pdf()
    return pdf_bytes, addendum["addendum_no"]


async def generate_teacher_contract_pdf(contract_id: str) -> tuple[bytes, str] | None:
    """產生教師合約 PDF，回傳 (pdf_bytes, contract_no) 或 None"""
    data = await _fetch_teacher_contract_data(contract_id)
    if not data:
        return None

    template = env.get_template("teacher_contract.html")
    html_str = template.render(
        contract=data["contract"],
        details=data["details"],
        work_schedules=data["work_schedules"],
        status_label=STATUS_LABELS.get(data["contract"]["contract_status"], data["contract"]["contract_status"]),
        employment_type_label=EMPLOYMENT_TYPE_LABELS.get(data["contract"].get("employment_type", ""), "-"),
        detail_type_labels=TEACHER_DETAIL_TYPE_LABELS,
        weekday_labels=WEEKDAY_LABELS,
    )

    pdf_bytes = HTML(string=html_str).write_pdf()
    return pdf_bytes, data["contract"]["contract_no"]


# ============================================
# DOCX 合約產生（基於 2026 新 EOP 課程合約範本）
# ============================================

def _replace_in_paragraph(paragraph, replacements: dict):
    """替換段落中的佔位符，保留格式"""
    for run in paragraph.runs:
        for key, value in replacements.items():
            if key in run.text:
                run.text = run.text.replace(key, str(value))


def _replace_in_cell(cell, replacements: dict):
    """替換表格儲存格中的佔位符"""
    for paragraph in cell.paragraphs:
        _replace_in_paragraph(paragraph, replacements)


async def generate_student_contract_docx(contract_id: str) -> tuple[bytes, str] | None:
    """
    從 docx 範本產生學生合約文件。
    回傳 (docx_bytes, filename) 或 None。
    """
    data = await _fetch_student_contract_data(contract_id)
    if not data:
        return None

    contract = data["contract"]
    details = data["details"]

    # 計算欄位值
    student_name = contract.get("student_name", "")
    total_lessons = contract.get("total_lessons", 0)
    total_amount = contract.get("total_amount", 0)

    # 學生聯絡資訊
    student_info = {"phone": "", "email": ""}
    if contract.get("student_id"):
        student_rows = await supabase_service.table_select(
            table="students", select="phone,email",
            filters={"id": contract["student_id"]},
        )
        if student_rows:
            student_info = student_rows[0]

    # 贈送堂數（compensation 類型的明細）
    bonus_lessons = 0
    for d in details:
        if d.get("detail_type") == "compensation":
            bonus_lessons += int(d.get("amount", 0))

    # 購買堂數 = 總堂數 - 贈送堂數
    purchased_lessons = total_lessons - bonus_lessons

    # 課程名稱 + 時長（從 lesson_price 明細取第一筆）
    course_name = ""
    duration_minutes = 50  # 預設
    for d in details:
        if d.get("detail_type") == "lesson_price" and d.get("course_name"):
            course_name = d["course_name"]
            if d.get("course_id"):
                course_rows = await supabase_service.table_select(
                    table="courses", select="duration_minutes",
                    filters={"id": d["course_id"]},
                )
                if course_rows and course_rows[0].get("duration_minutes"):
                    duration_minutes = course_rows[0]["duration_minutes"]
            break

    # 日期計算
    start_date = contract.get("start_date", "")
    end_date = contract.get("end_date", "")
    if isinstance(start_date, date):
        start_date_str = start_date.strftime("%Y/%m/%d")
    else:
        start_date_str = str(start_date).replace("-", "/") if start_date else ""

    # 期限月數
    period_months = 0
    if start_date and end_date:
        sd = start_date if isinstance(start_date, date) else date.fromisoformat(str(start_date))
        ed = end_date if isinstance(end_date, date) else date.fromisoformat(str(end_date))
        period_months = max(1, round((ed - sd).days / 30))

    # 請假額度
    leave_quota = math.ceil(total_lessons * 0.2) if total_lessons else 0

    # 金額格式化
    if total_amount:
        total_amount_str = f"{int(total_amount):,}"
    else:
        total_amount_str = "0"

    # 簽約日期
    sign_date = date.today().strftime("%Y/%m/%d")

    # 佔位符對照
    replacements = {
        "{{course_name}}": course_name or "（課程名稱）",
        "{{purchased_lessons}}": str(purchased_lessons),
        "{{duration_minutes}}": str(duration_minutes),
        "{{bonus_lessons}}": str(bonus_lessons),
        "{{total_amount}}": total_amount_str,
        "{{start_date}}": start_date_str or "____/____/____",
        "{{all_lessons}}": str(total_lessons),
        "{{period_months}}": str(period_months),
        "{{leave_quota}}": str(leave_quota),
        "{{sign_date}}": sign_date,
        "{{student_name}}": student_name or "",  # 用於合約內文（非簽署欄）
        "{{student_phone}}": student_info.get("phone", "") or "",
        "{{student_email}}": student_info.get("email", "") or "",
    }

    # 讀取範本並替換
    template_path = os.path.join(TEMPLATE_DIR, "contract_template.docx")
    doc = DocxDocument(template_path)

    for paragraph in doc.paragraphs:
        if "{{" in paragraph.text:
            _replace_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "{{" in cell.text:
                    _replace_in_cell(cell, replacements)

    # 輸出 bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    contract_no = contract.get("contract_no", "contract")
    filename = f"{contract_no}_{student_name}_合約.docx"

    return buffer.getvalue(), filename
